import csv
import docx

INPUT_FILES_PATH = r'/home/mouadziani/Documents/workspace/python/excel_files/'
OUTPUT_FILES_PATH = r'/home/mouadziani/Documents/workspace/python/doc_files/'

# Helpers 
def load_data_from_csv(file_name):
    file = open(INPUT_FILES_PATH + file_name, "r")
    reader = csv.reader(file)
    lines = []
    for line in reader:
        line = [line[0], clean_case(line[1])]
        lines.append(line)
    return lines[2:]


def clean_case(item):
    return (item.replace(',', '')).strip()


def save_array_as_doc_file(file_name, array):
    doc = docx.Document()
    doc.add_paragraph('================================================================================')
    doc.add_paragraph('=========================== Excel diff checker V1 ==============================')
    doc.add_paragraph('================================================================================')
    for line in array:
        doc.add_paragraph(str(line))
    doc.save(OUTPUT_FILES_PATH + file_name)


def check_difference(array_1, array_2):
    check_for = array_1 if len(array_1) > len(array_2) else array_2
    check_in = array_1 if len(array_1) < len(array_2) else array_2
    check_in_file_name = 'file_1.csv' if len(array_1) < len(array_2) else 'file_1.csv'
    outpu_messages = []
    for item_check_for in check_for:
        exists = False
        for item_check_in in check_in:
            if item_check_for[0] == item_check_in[0]:
                exists = True
                if item_check_for[1] != item_check_in[1]:
                    msg = 'There\'s a difference in (' + str(item_check_for[0]) + ') | File 1 : ' + str(item_check_for[1]) + ', File 2: ' + str(item_check_in[1])
                    outpu_messages.append(msg)
                    print(msg)
                    msg = '------------------------------------------------------------------------------'
                    outpu_messages.append(msg)
                    print(msg)
                    break
        if not exists:
            msg = 'The item (' + str(item_check_for[0]) + ') has not been not found in the [' + check_in_file_name + ']'
            outpu_messages.append(msg)
            print(msg)
            msg = '------------------------------------------------------------------------------'
            outpu_messages.append(msg)
            print(msg)
        save_array_as_doc_file('result_diff.docx', outpu_messages)

print('\n================================================================================')
print('=========================== Excel diff checker V1 ==============================')
print('================================================================================\n')
data_file_1 = load_data_from_csv('file_1.csv')
data_file_2 = load_data_from_csv('file_2.csv')
check_difference(data_file_1, data_file_2)